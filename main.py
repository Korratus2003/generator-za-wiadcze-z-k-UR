import sys
import pandas as pd
import os
from docx import Document
from datetime import datetime
import argparse
from docx.oxml import OxmlElement

template_file = 'szablon.docx'
output_dir = 'zaswiadczenia'

def parse_name_email(email, indeks):
    base = email.split('@')[0]
    rest = base.replace(str(indeks), '')
    if len(rest) >= 2:
        return rest[0].upper(), rest[1].upper()
    else:
        return 'X', 'Y'

def extract_date(date_str):
    text = str(date_str).strip().lstrip('/')
    parts = text.split()
    if len(parts) == 0:
        return None
    date_part = parts[0]
    for fmt in ('%m/%d/%Y', '%d/%m/%Y'):
        try:
            dt = datetime.strptime(date_part, fmt)
            return dt
        except ValueError:
            continue
    return None

def generate_certificate(row):
    email = row['Adres e-mail']
    indeks = ''.join([s for s in email if s.isdigit()][0:6])
    data_wypelnienia = extract_date(row['Godzina rozpoczęcia'])
    if data_wypelnienia is None:
        print(f"Niepoprawna data dla wiersza: {row['Id']}")
        return
    kierunek = row['Kierunek']
    osiagniecia = [o.strip() for o in str(row['osiągnięcia']).split(';') if o.strip()]
    od_kiedy = extract_date(row['Od kiedy jesteś w kole'])
    od_kiedy_str = od_kiedy.strftime('%d.%m.%Y') if od_kiedy else '??.??.????'

    do_kiedy = datetime.now()
    do_kiedy_str = do_kiedy.strftime('%d.%m.%Y')
    imie_nazwisko = row['Nazwa']

    data_wypelnienia_fmt = data_wypelnienia.strftime('%d.%m.%Y')

    doc = Document(template_file)

    # Podstawianie prostych placeholderów
    for para in doc.paragraphs:
        if '{IMIE_NAZWISKO}' in para.text:
            para.text = para.text.replace('{IMIE_NAZWISKO}', imie_nazwisko)
        if '{NUMER_ALBUMU}' in para.text:
            para.text = para.text.replace('{NUMER_ALBUMU}', indeks)
        if '{KIERUNEK}' in para.text:
            para.text = para.text.replace('{KIERUNEK}', kierunek)
        if '{DATA}' in para.text:
            para.text = para.text.replace('{DATA}', data_wypelnienia_fmt)
        if '{OD_KIEDY}' in para.text:
            para.text = para.text.replace('{OD_KIEDY}', od_kiedy_str)
        if '{DO_KIEDY}' in para.text:
            para.text = para.text.replace('{DO_KIEDY}', do_kiedy_str)

    # Wstawianie listy osiągnięć w miejscu {OSIAGNIECIA}
    for para in doc.paragraphs:
        if '{OSIAGNIECIA}' in para.text:
            parent = para._element.getparent()
            index = parent.index(para._element)

            # Usuń placeholder z paragrafu (można też usunąć cały paragraf)
            para.text = para.text.replace('{OSIAGNIECIA}', '').strip()

            # Wstaw paragrafy z osiągnięciami dokładnie po tym paragrafie
            for idx, os_text in enumerate(osiagniecia):
                new_p = OxmlElement('w:p')
                new_r = OxmlElement('w:r')
                new_t = OxmlElement('w:t')
                new_t.text = f"{idx + 1}. {os_text}"

                new_r.append(new_t)
                new_p.append(new_r)

                parent.insert(index + 1 + idx, new_p)
            break

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{imie_nazwisko} zaświadczenie.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"Wygenerowano {filepath}")

def main():
    parser = argparse.ArgumentParser(description='Generator zaświadczeń ze ścieżką do pliku CSV.')
    parser.add_argument('csv_file', help='Ścieżka do pliku CSV z danymi')

    if len(sys.argv) == 1:
        parser.print_help()
        print("\nBłąd: Brak wymaganej ścieżki do pliku CSV. Proszę podać plik jako argument.")
        sys.exit(1)

    args = parser.parse_args()

    if not os.path.isfile(args.csv_file):
        print(f"Błąd: plik {args.csv_file} nie istnieje.")
        sys.exit(1)

    df = pd.read_csv(args.csv_file, encoding='utf-8')

    for _, row in df.iterrows():
        generate_certificate(row)

if __name__ == '__main__':
    main()
